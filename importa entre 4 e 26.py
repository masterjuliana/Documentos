import pandas as pd
import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from collections import defaultdict
import re

# ==============================================================================
# 1. CONFIGURAÇÃO INICIAL
# ==============================================================================

# Configuração do logger para registrar mensagens em um arquivo e no console
logging.basicConfig(
    level=logging.INFO, # Define o nível mínimo de mensagens a serem registradas (INFO, WARNING, ERROR, DEBUG)
    format="%(asctime)s - %(levelname)s - %(message)s", # Formato da mensagem de log
    handlers=[
        logging.FileHandler("processamento.log", encoding="utf-8"), # Salva logs em um arquivo
        logging.StreamHandler() # Exibe logs no console
    ]
)

# Definição dos caminhos dos arquivos usando Pathlib para compatibilidade entre sistemas
BASE_DIR = Path(__file__).parent # Diretório base onde o script está sendo executado
ARQUIVO_EXCEL = BASE_DIR / "Consideracoes-sobre-a-Consulta_Publica_Decreto_7217.2010.xlsx" # Caminho para o arquivo Excel de entrada
ARQUIVO_WORD_ENTRADA = BASE_DIR / "saida.docx" # Caminho para o documento Word de entrada (modelo)
ARQUIVO_WORD_SAIDA = BASE_DIR / "saida_4-26.docx" # Caminho para o documento Word de saída (com as alterações)

# --- NOVAS CONFIGURAÇÕES DE FILTRO E ESTILO ---
# Define os critérios para filtrar os itens do Excel a serem processados.
# As opções são:
# - 'all': processa todos os itens numéricos válidos.
# - 'zero_and_greater_than_X': processa o item 0 E itens maiores que X (onde X é o número definido abaixo).
# - [lista_de_itens]: processa apenas os itens listados (ex: [1, 5, 10]).
FILTRO_ITEM_EXCEL = list(range(4, 27)) # <--- **ESSA É A ALTERAÇÃO CRUCIAL PARA FILTRAR DO 4 AO 26**
VALOR_REFERENCIA_FILTRO = 52 # Este valor não será usado com a configuração acima, mas pode ser mantido.

# Configurações de estilo para o texto das contribuições no Word
ESTILOS_CONTRIBUICAO = {
    'fonte': 'Calibri',
    'tamanho': 8,
    'cor': 'auto', # 'auto' para cor padrão, ou um código de cor (ex: '0000FF' para azul)
}


# ==============================================================================
# 2. FUNÇÕES AUXILIARES
# ==============================================================================

def adicionar_linha_separadora(paragraph):
    """
    Adiciona uma borda inferior a um parágrafo no documento Word, criando uma linha separadora visual.
    """
    p_pr = paragraph._p.get_or_add_pPr() # Obtém ou adiciona propriedades de parágrafo
    p_bdr = OxmlElement('w:pBdr') # Cria um elemento XML para bordas de parágrafo
    p_pr.append(p_bdr) # Adiciona o elemento de borda às propriedades do parágrafo
    
    bottom_border = OxmlElement('w:bottom') # Cria um elemento XML para a borda inferior
    bottom_border.set(qn('w:val'), 'single') # Estilo da linha: 'single' (linha única)
    bottom_border.set(qn('w:sz'), '6') # Tamanho da linha (em oitavos de ponto, 6 = 0.75pt)
    bottom_border.set(qn('w:space'), '1') # Espaço entre a linha e o texto
    bottom_border.set(qn('w:color'), 'auto') # Cor da linha (automática)
    
    p_bdr.append(bottom_border) # Adiciona a borda inferior ao elemento de bordas do parágrafo
    logging.debug("Linha separadora adicionada.")

def clear_cell(cell):
    """
    Remove todo o conteúdo de uma célula do Word (parágrafos e tabelas aninhadas).
    """
    for paragraph in cell.paragraphs:
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    for table in cell.tables:
        t_element = table._element
        t_element.getparent().remove(t_element)
    # Adiciona um parágrafo vazio para garantir que a célula não fique completamente vazia
    # o que pode causar problemas de layout em algumas versões do Word
    cell.add_paragraph()
    logging.debug("Conteúdo da célula limpo.")

def formatar_celula_com_contribuicoes(celula, contribuicoes: list, estilos: dict):
    """
    Preenche uma célula do documento Word com uma lista de contribuições, formatando cada uma
    e adicionando linhas separadoras entre elas, usando um dicionário de estilos.

    Args:
        celula: O objeto Cell do python-docx a ser preenchido.
        contribuicoes (list): Lista de dicionários, onde cada dicionário representa uma contribuição
                               com chaves como 'numero', 'Titulo da Contribuição', 'texto' e 'nome'.
        estilos (dict): Dicionário contendo as configurações de estilo (fonte, tamanho, cor).
    """
    fonte = estilos.get('fonte', 'Calibri')
    tamanho = estilos.get('tamanho', 8)
    # cor = estilos.get('cor', '000000') # Para uso futuro se precisar de cor de fonte específica

    # Limpa o conteúdo existente da célula antes de adicionar o novo
    clear_cell(celula)
    
    for idx, contrib in enumerate(contribuicoes):
        # Adicionar número da contribuição
        p_num = celula.add_paragraph() # Cria um novo parágrafo na célula
        run_num = p_num.add_run(contrib['numero']) # Adiciona o número da contribuição
        run_num.font.name = fonte
        run_num.font.size = Pt(tamanho)
        run_num.bold = True # Número em negrito
        logging.debug(f"   Adicionado número da contribuição: {contrib['numero']}")
        
        # Adicionar Titulo da Contribuição 
        p_titulo = celula.add_paragraph()
        run_titulo = p_titulo.add_run(contrib['Titulo da Contribuição'])
        run_titulo.font.name = fonte
        run_titulo.font.size = Pt(tamanho)
        logging.debug(f"   Adicionado título da contribuição: {contrib['Titulo da Contribuição']}")
        
        # Adicionar texto da contribuição
        p_texto = celula.add_paragraph(contrib['texto'])
        for run in p_texto.runs: # Itera sobre os runs do parágrafo para aplicar a formatação
            run.font.name = fonte
            run.font.size = Pt(tamanho)
        logging.debug(f"   Adicionado texto da contribuição (início): {contrib['texto'][:50]}...")
        
        # Adicionar nome do autor
        p_autor = celula.add_paragraph()
        run_autor = p_autor.add_run(f"({contrib['nome']})")
        run_autor.font.name = fonte
        run_autor.font.size = Pt(tamanho)
        run_autor.italic = True # Nome do autor em itálico
        logging.debug(f"   Adicionado nome do autor: {contrib['nome']}")
        
        # Adicionar linha separadora (exceto após a última contribuição)
        if idx < len(contribuicoes) - 1:
            p_separador = celula.add_paragraph()
            adicionar_linha_separadora(p_separador)
            logging.debug("   Adicionada linha separadora entre contribuições.")
    logging.debug("Célula formatada com todas as contribuições.")

def validar_item(item: str) -> int:
    """
    Valida e converte valores de item para um número inteiro, removendo caracteres não numéricos.
    Retorna None se a string resultante estiver vazia ou se houver erro na conversão.
    """
    try:
        # Usa expressão regular para remover todos os caracteres que não são dígitos
        item_limpo = re.sub(r'\D', '', str(item))
        if item_limpo: # Se a string não estiver vazia após a limpeza
            logging.debug(f"Validando item: '{item}' -> Limpo: '{item_limpo}' -> Convertido: {int(item_limpo)}")
            return int(item_limpo) # Converte para inteiro
        logging.debug(f"Validando item: '{item}' -> Limpo: '{item_limpo}' (vazio), retornando None.")
        return None # Retorna None se a string limpa estiver vazia
    except (ValueError, TypeError) as e:
        # Captura erros de conversão ou tipo inválido, e retorna None
        logging.warning(f"Falha ao validar item '{item}': {e}. Retornando None.")
        return None

# ==============================================================================
# 3. FUNÇÃO PRINCIPAL DE EXECUÇÃO
# ==============================================================================

def main() -> None:
    """
    Função principal que orquestra a leitura do Excel, processamento dos dados
    e atualização do documento Word.
    """
    logging.info(f"Iniciando processamento. Caminho do Excel: {ARQUIVO_EXCEL}")
    
    # --- Leitura e preparação dos dados do Excel ---
    try:
        # Lê o arquivo Excel, especificando o motor, tipo de dados e colunas a serem usadas
        df = pd.read_excel(
            ARQUIVO_EXCEL,
            engine="openpyxl",
            dtype=str, # Lê todas as colunas como string para evitar problemas de tipo
            usecols=['Item CP alterado', 'Numero', 'Titulo da Contribuição ', 'Texto', 'Nome']
        )
        logging.info(f"Excel lido com sucesso. Colunas: {df.columns.tolist()}")
        logging.debug(f"DataFrame lido (primeiras 5 linhas):\n{df.head().to_string()}")
    except FileNotFoundError:
        logging.error(f"Erro: Arquivo Excel não encontrado em '{ARQUIVO_EXCEL}'. Verifique o caminho.")
        return
    except KeyError as e:
        logging.error(f"Erro: Coluna '{e}' não encontrada no arquivo Excel. Verifique os nomes das colunas.")
        return
    except Exception as e:
        logging.exception(f"Falha inesperada na leitura do Excel: {e}")
        return

    # --- Pré-processamento de dados do DataFrame ---
    for col in df.columns:
        if df[col].dtype == 'object':
            df[col] = df[col].astype(str).str.replace('_x000D_', '\n', regex=False)
    df.columns = df.columns.str.strip()
    logging.info("Pré-processamento de colunas e dados do Excel concluído.")
    logging.debug(f"DataFrame após pré-processamento (primeiras 5 linhas):\n{df.head().to_string()}")
    
    # --- Filtragem e validação de itens ---
    df['Item_validado'] = df['Item CP alterado'].apply(validar_item)
    logging.info("Coluna 'Item_validado' criada.")
    
    df = df.dropna(subset=['Item_validado'])
    logging.info(f"Removidas linhas com 'Item_validado' inválido. Restam {len(df)} linhas.")

    df['Item_validado'] = df['Item_validado'].astype(int)
    logging.info("Coluna 'Item_validado' convertida para inteiro.")

    # --- FILTRAGEM DE ACORDO COM A NOVA CONFIGURAÇÃO ---
    df_para_processar = pd.DataFrame() # Inicializa um DataFrame vazio para os dados filtrados

    if FILTRO_ITEM_EXCEL == 'all':
        df_para_processar = df.copy()
        logging.info(f"Configurado para processar TODOS os itens válidos. Total: {len(df_para_processar)} linhas.")
    elif isinstance(FILTRO_ITEM_EXCEL, list): # Esta é a condição que será usada agora!
        # Filtra por uma lista específica de itens
        df_para_processar = df[df['Item_validado'].isin(FILTRO_ITEM_EXCEL)].copy()
        logging.info(f"Configurado para processar itens específicos: {FILTRO_ITEM_EXCEL}. Total: {len(df_para_processar)} linhas.")
    elif FILTRO_ITEM_EXCEL == 'zero_and_greater_than_52': # Condição original, não será ativada com a nova FILTRO_ITEM_EXCEL
        # Filtra o item 0 OU itens maiores que VALOR_REFERENCIA_FILTRO
        df_para_processar = df[(df['Item_validado'] == 0) | (df['Item_validado'] > VALOR_REFERENCIA_FILTRO)].copy()
        logging.info(f"Configurado para processar item 0 E itens maiores que {VALOR_REFERENCIA_FILTRO}. Total: {len(df_para_processar)} linhas.")
    # A condição `FILTRO_ITEM_EXCEL == '4'` que você tinha antes, foi removida daqui,
    # pois a forma correta de filtrar uma lista de itens é com `isinstance(FILTRO_ITEM_EXCEL, list)`.
    else:
        logging.warning(f"Configuração de FILTRO_ITEM_EXCEL inválida: '{FILTRO_ITEM_EXCEL}'. Nenhum item será processado.")
        return # Encerra a execução se a configuração for inválida

    if df_para_processar.empty:
        logging.warning("Nenhum dado encontrado para processamento após a filtragem! Verifique o Excel ou as configurações de filtro.")
        return # Encerra a execução se não houver dados para o filtro

    # A partir daqui, usaremos o DataFrame filtrado conforme a configuração
    df = df_para_processar

    # --- Agrupamento de contribuições em formato estruturado ---
    contrib_por_item = defaultdict(list)
    df = df.sort_values(by=["Item_validado", "Numero"])
    
    for _, row in df.iterrows():
        item = row['Item_validado']
        contrib = {
            'numero': f"{row['Numero']}" if pd.notna(row['Numero']) else "[sem número]",
            'Titulo da Contribuição': row['Titulo da Contribuição'] if pd.notna(row['Titulo da Contribuição']) else "[sem Titulo da Contribuição]",
            'texto': row['Texto'] if pd.notna(row['Texto']) else "[sem texto]",
            'nome': row['Nome'] if pd.notna(row['Nome']) else "[autor desconhecido]"
        }
        contrib_por_item[item].append(contrib)
    
    logging.info(f"Contribuições agrupadas para {len(contrib_por_item)} itens únicos encontrados no Excel (após filtro).")
    logging.debug(f"Itens agrupados: {list(contrib_por_item.keys())}") # Exibe os itens que foram agrupados

    # --- Processamento do documento Word ---
    try:
        doc = Document(ARQUIVO_WORD_ENTRADA)
        logging.info(f"Documento Word de entrada '{ARQUIVO_WORD_ENTRADA}' aberto com sucesso.")
    except FileNotFoundError:
        logging.error(f"Erro: Arquivo Word não encontrado em '{ARQUIVO_WORD_ENTRADA}'. Verifique o caminho.")
        return
    except Exception as e:
        logging.exception(f"Falha inesperada ao abrir documento Word: {e}")
        return

    itens_atualizados = 0
    
    # Itera sobre todas as tabelas e linhas do documento Word
    for idx_tabela, tabela in enumerate(doc.tables):
        logging.info(f"Processando Tabela {idx_tabela + 1} no Word...")
        for idx_linha, linha in enumerate(tabela.rows):
            if len(linha.cells) < 3:
                logging.debug(f"   Linha {idx_linha + 1} da Tabela {idx_tabela + 1}: Ignorada (menos de 3 células).")
                continue
            
            celula_item = linha.cells[0]
            item_cp_text = celula_item.text.strip()
            item_cp = validar_item(item_cp_text)
            
            logging.info(f"   Linha {idx_linha + 1} da Tabela {idx_tabela + 1}: Texto da célula de item '{item_cp_text}' -> Validado como: {item_cp}")
            
            if item_cp is not None and item_cp in contrib_por_item:
                logging.info(f"     *** Item {item_cp} encontrado no Word e com contribuições no Excel. ATUALIZANDO CÉLULA... ***")
                formatar_celula_com_contribuicoes(
                    linha.cells[2], # A célula a ser preenchida (terceira coluna)
                    contrib_por_item[item_cp], # As contribuições agrupadas para este item
                    estilos=ESTILOS_CONTRIBUICAO # Passa o dicionário de estilos
                )
                itens_atualizados += 1
            else:
                logging.info(f"     Item {item_cp} (ou inválido/nulo) NÃO possui contribuições no Excel ou não foi encontrado para atualização (com base no filtro ativo).")

    # --- Salvamento do documento Word ---
    try:
        doc.save(ARQUIVO_WORD_SAIDA)
        logging.info(f"Documento Word salvo com sucesso: {ARQUIVO_WORD_SAIDA}")
        logging.info(f"Total de itens atualizados: {itens_atualizados}")
    except Exception as e:
        logging.exception(f"Erro ao salvar documento Word: {e}")

# ==============================================================================
# 4. EXECUÇÃO PRINCIPAL
# ==============================================================================

if __name__ == "__main__":
    main()