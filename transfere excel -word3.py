import pandas as pd
import logging
from docx import Document
from docx.shared import Pt

# ==============================================================================
# 1. CONFIGURAÇÃO INICIAL
# ==============================================================================

# Configura o sistema de logging para fornecer feedback detalhado durante a execução.
logging.basicConfig(
    level=logging.INFO, 
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Definição dos caminhos dos ficheiros de entrada e saída.
ARQUIVO_EXCEL = r"W:\MINISTÉRIO DAS CIDADES\Consulta 7217\python\Consideracoes-sobre-a-Consulta_Publica_Decreto_7217.2010.xlsx"
ARQUIVO_WORD_ENTRADA = r"W:\MINISTÉRIO DAS CIDADES\Consulta 7217\python\saida.docx"
ARQUIVO_WORD_SAIDA = r"W:\MINISTÉRIO DAS CIDADES\Consulta 7217\python\saida_ajustada.docx"


# ==============================================================================
# 2. FUNÇÃO AUXILIAR PARA FORMATAÇÃO
# ==============================================================================

def formatar_celula(celula, texto, nome_fonte='Calibri', tamanho_fonte=10, negrito=False):
    """
    Aplica texto e formatação a uma célula de uma tabela do Word.

    Args:
        celula (obj): O objeto da célula a ser formatado.
        texto (str): O conteúdo a ser inserido na célula.
        nome_fonte (str): O nome da fonte a ser aplicada.
        tamanho_fonte (int): O tamanho da fonte em pontos (Pt).
        negrito (bool): Se o texto deve ou não estar em negrito.
    """
    celula.text = texto
    run = celula.paragraphs[0].runs[0]
    run.font.name = nome_fonte
    run.font.size = Pt(tamanho_fonte)
    run.bold = negrito


# ==============================================================================
# 3. FUNÇÃO PRINCIPAL DE EXECUÇÃO
# ==============================================================================

def main():
    """
    Função principal que orquestra a leitura, processamento e escrita dos dados.
    """
    # --- Leitura e Validação do Excel ---
    logging.info(f"Iniciando a leitura do arquivo Excel: {ARQUIVO_EXCEL}")
    try:
        df = pd.read_excel(ARQUIVO_EXCEL, engine="openpyxl")
    except FileNotFoundError:
        logging.error(f"O arquivo '{ARQUIVO_EXCEL}' não foi encontrado. O programa será encerrado.")
        raise
    except Exception as e:
        logging.error(f"Ocorreu um erro inesperado ao ler o arquivo Excel: {e}")
        raise

    logging.info("Arquivo Excel lido com sucesso.")

    # --- Limpeza e Preparação dos Dados ---
    df.columns = df.columns.str.strip()
    colunas_desejadas = ['Item CP alterado', 'Numero', 'Titulo da Contribuição', 'Texto', 'Justificativa', 'Nome']
    
    colunas_ausentes = [col for col in colunas_desejadas if col not in df.columns]
    if colunas_ausentes:
        raise ValueError(f"Colunas ausentes no Excel: {', '.join(colunas_ausentes)}")

    df = df[colunas_desejadas]
    if df.empty:
        raise ValueError("Nenhum dado encontrado após a seleção das colunas!")

    df = df.replace(r'_x000D_', '"', regex=True)

    # --- Transformação e Filtragem dos Dados ---
    df["Item CP alterado"] = pd.to_numeric(df["Item CP alterado"], errors="coerce")
    df = df.dropna(subset=["Item CP alterado"])
    df["Item CP alterado"] = df["Item CP alterado"].astype(int)
    df = df[df["Item CP alterado"] >= 100]
    
    logging.info("Dados limpos, transformados e filtrados.")

    # --- Agrupamento das Contribuições ---
    df_agrupado = df.groupby("Item CP alterado").agg({
        "Numero": lambda x: ", ".join(x.astype(str)),
        "Texto": lambda x: "\n".join(x.astype(str)),
        "Justificativa": lambda x: "\n".join(x.astype(str)),
        "Nome": lambda x: ", ".join(x.astype(str))
    }).reset_index()

    logging.info("Contribuições agrupadas por item com sucesso.")

    # --- Atualização do Documento Word ---
    logging.info(f"Abrindo o documento Word: {ARQUIVO_WORD_ENTRADA}")
    doc = Document(ARQUIVO_WORD_ENTRADA)

    itens_atualizados = 0
    for table in doc.tables:
        for row in table.rows:
            item_cp_cell_text = row.cells[0].text.strip()

            if item_cp_cell_text.isdigit():
                item_cp_word = int(item_cp_cell_text)
                
                matching_row = df_agrupado[df_agrupado["Item CP alterado"] == item_cp_word]

                if not matching_row.empty:
                    logging.info(f"Encontrada correspondência para o item '{item_cp_word}'. Atualizando a tabela...")
                    
                    # --- INÍCIO DA ALTERAÇÃO ---
                    
                    # 1. Extrair todos os valores da linha correspondente
                    item_val = str(matching_row["Item CP alterado"].values[0])
                    numero_val = str(matching_row["Numero"].values[0])
                    texto_val = str(matching_row['Texto'].values[0])
                    justificativa_val = str(matching_row['Justificativa'].values[0])
                    nome_val = str(matching_row['Nome'].values[0])

                    # 2. Construir a nova string para a coluna "Visões", incluindo o Número
                    texto_completo_com_numero = (
                        f"Contribuição(ões) N.º: {numero_val}\n\n"
                        f"{texto_val}\n"
                        f"{justificativa_val}\n"
                        f"({nome_val})"
                    )

                    # 3. Utiliza a função auxiliar para formatar cada célula
                    # Célula 0: Item (permanece igual)
                    formatar_celula(row.cells[0], texto=item_val, tamanho_fonte=10, negrito=True)
                    
                    # Célula 1: Número (agora fica em branco, pois o valor foi movido)
                    formatar_celula(row.cells[1], texto="") 
                    
                    # Célula 2: Visões (agora inclui o número no seu conteúdo)
                    formatar_celula(row.cells[2], texto=texto_completo_com_numero, tamanho_fonte=8)
                    
                    # --- FIM DA ALTERAÇÃO ---
                    
                    itens_atualizados += 1

    logging.info(f"Processamento do Word concluído. {itens_atualizados} itens foram atualizados na tabela.")

    # --- Gravação do Ficheiro Final ---
    doc.save(ARQUIVO_WORD_SAIDA)
    logging.info(f"Arquivo Word final salvo com sucesso em: {ARQUIVO_WORD_SAIDA}")


# ==============================================================================
# 4. PONTO DE ENTRADA DO SCRIPT
# ==============================================================================

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        logging.error(f"Ocorreu um erro fatal durante a execução do script: {e}")