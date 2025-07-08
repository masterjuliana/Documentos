import matplotlib.pyplot as plt
import re
import os
import PyPDF2
from openpyxl import Workbook
from transformers import pipeline
import pandas as pd
from docx import Document

# Recriar os dados após reset
data = {
    "Município": [
        "Rio de Janeiro", "Niterói", "Mesquita", "Belford Roxo", "Duque de Caxias",
        "Nova Iguaçu", "Queimados", "Japeri", "Magé", "São Gonçalo", "Maricá", "Petrópolis"
    ],
    "Coleta de Esgoto (%) [SNIS 2021]": [100, 100, 100, 100, 89, 91, 91, 66, 20, 90, 37, 97],
    "Tratamento (%) [SEAS 2022]": [27.61, 96.47, 0, 0, 0, 0, 18.06, 0, 0, 0, 3.31, 47.65],
    "Reúso Estimado (%)": [2, 10, 0, 0, 0, 0, 1.5, 0, 0, 0, 0.5, 4]
}
df = pd.DataFrame(data)
df.set_index("Município", inplace=True)

# Criar documento Word
from docx.shared import Inches
doc = Document()
doc.add_heading('Relatório Técnico: Panorama de Saneamento e Potencial de Reúso – Região Metropolitana do RJ', 0)

# Introdução
doc.add_paragraph(
    "Este relatório apresenta uma análise comparativa entre os índices de coleta de esgoto, "
    "tratamento efetivo e o potencial estimado de reúso de efluentes em municípios da Região "
    "Metropolitana do Rio de Janeiro. Os dados utilizados foram extraídos dos relatórios do "
    "SNIS 2021 e SEAS-RJ 2022, com estimativas de reúso fundamentadas em padrões de cobertura "
    "e infraestrutura identificados em cada localidade."
)

# Tabela de dados
doc.add_heading('Tabela Comparativa', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Município'
hdr_cells[1].text = 'Coleta de Esgoto (%) [SNIS 2021]'
hdr_cells[2].text = 'Tratamento (%) [SEAS 2022]'
hdr_cells[3].text = 'Reúso Estimado (%)'

for index, row in df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(index)
    row_cells[1].text = f"{row['Coleta de Esgoto (%) [SNIS 2021]']}%"
    row_cells[2].text = f"{row['Tratamento (%) [SEAS 2022]']}%"
    row_cells[3].text = f"{row['Reúso Estimado (%)']}%"

# Análise
doc.add_heading('Análise Técnica', level=1)
doc.add_paragraph(
    "A análise revela que, embora a maioria dos municípios apresente altas taxas de coleta de esgoto "
    "(com destaque para Niterói, Mesquita, Belford Roxo e Rio de Janeiro), a proporção de esgoto que "
    "recebe tratamento efetivo é drasticamente menor. Em cidades como Duque de Caxias, Nova Iguaçu, "
    "Mesquita e Belford Roxo, a taxa de tratamento é nula, indicando despejo direto nos corpos hídricos.\n\n"
    "O reúso de efluentes aparece como solução estratégica para aumentar a resiliência hídrica e reduzir "
    "a poluição em áreas urbanas densamente povoadas. Municípios como Niterói e Petrópolis, que já possuem "
    "infraestrutura de tratamento consolidada, apresentam maior viabilidade imediata para projetos de reúso."
)

# Conclusão
doc.add_heading('Conclusão', level=1)
doc.add_paragraph(
    "Os dados reforçam a necessidade de investimentos urgentes na expansão do tratamento de esgoto e na "
    "adoção de políticas de reúso, sobretudo em áreas críticas como a Baixada Fluminense. A integração entre "
    "saneamento, planejamento urbano e soluções circulares como o reúso de efluentes deve ser central em "
    "estratégias futuras de sustentabilidade e justiça ambiental."
)

# Salvar documento
file_path = "/mnt/data/Relatorio_Tecnico_Saneamento_Reuso_RJ.docx"
doc.save(file_path)
file_path




# Gerar gráfico de barras para Coleta, Tratamento e Reúso
fig, ax = plt.subplots(figsize=(12, 6))
bar_width = 0.25
index = range(len(df))

# Barras
plt.bar(index, df["Coleta de Esgoto (%) [SNIS 2021]"], bar_width, label="Coleta (%)")
plt.bar([i + bar_width for i in index], df["Tratamento (%) [SEAS 2022]"], bar_width, label="Tratamento (%)")
plt.bar([i + 2 * bar_width for i in index], df["Reúso Estimado (%)"], bar_width, label="Reúso (%)")

# Rótulos
plt.xlabel('Municípios')
plt.ylabel('Percentual (%)')
plt.title('Comparativo de Coleta, Tratamento e Reúso de Esgoto')
plt.xticks([i + bar_width for i in index], df.index, rotation=45, ha='right')
plt.legend()
plt.tight_layout()

# Salvar gráfico como imagem
img_path = "/mnt/data/Grafico_Saneamento_Reuso.png"
plt.savefig(img_path)
plt.close()

img_path
