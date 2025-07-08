import pandas as pd
import logging
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docx.shared import Pt
import os
import threading

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def formatar_celula(celula, texto, nome_fonte='Calibri', tamanho_fonte=10, negrito=False):
    celula.text = texto
    run = celula.paragraphs[0].runs[0]
    run.font.name = nome_fonte
    run.font.size = Pt(tamanho_fonte)
    run.bold = negrito

def processar(arquivo_excel, arquivo_word_entrada, arquivo_word_saida, progresso_callback=None):
    try:
        df = pd.read_excel(arquivo_excel, engine="openpyxl")
        df.columns = df.columns.str.strip()
        df = df[['Item CP alterado', 'Numero', 'Titulo da Contribuição', 'Texto', 'Justificativa', 'Nome']]
        df = df.replace(r'_x000D_', '"', regex=True)
        df["Item CP alterado"] = pd.to_numeric(df["Item CP alterado"], errors="coerce")
        df = df.dropna(subset=["Item CP alterado"])
        df["Item CP alterado"] = df["Item CP alterado"].astype(int)
        df = df[df["Item CP alterado"] >= 100]

        df_agrupado = df.groupby("Item CP alterado").agg({
            "Numero": lambda x: ", ".join(x.astype(str)),
            "Texto": lambda x: "\n".join(x.astype(str)),
            "Justificativa": lambda x: "\n".join(x.astype(str)),
            "Nome": lambda x: ", ".join(x.astype(str))
        }).reset_index()

        doc = Document(arquivo_word_entrada)
        total_tabelas = len(doc.tables)
        itens_atualizados = 0

        for t_index, table in enumerate(doc.tables):
            for row in table.rows:
                if len(row.cells) < 3:
                    continue
                item_cp_cell_text = row.cells[0].text.strip()
                if item_cp_cell_text.isdigit():
                    item_cp_word = int(item_cp_cell_text)
                    matching_row = df_agrupado[df_agrupado["Item CP alterado"] == item_cp_word]
                    if not matching_row.empty:
                        item_val = str(matching_row["Item CP alterado"].values[0])
                        numero_val = str(matching_row["Numero"].values[0])
                        texto_val = str(matching_row['Texto'].values[0])
                        justificativa_val = str(matching_row['Justificativa'].values[0])
                        nome_val = str(matching_row['Nome'].values[0])
                        texto_completo_com_numero = (
                            f"Contribuição(ões) N.º: {numero_val}\n\n"
                            f"{texto_val}\n"
                            f"{justificativa_val}\n"
                            f"({nome_val})"
                        )
                        formatar_celula(row.cells[0], texto=item_val, tamanho_fonte=10, negrito=True)
                        formatar_celula(row.cells[1], texto="") 
                        formatar_celula(row.cells[2], texto=texto_completo_com_numero, tamanho_fonte=8)
                        itens_atualizados += 1

            # Atualiza progresso
            if progresso_callback:
                progresso_callback((t_index + 1) / total_tabelas * 100)

        doc.save(arquivo_word_saida)
        return f"Processamento concluído: {itens_atualizados} itens atualizados."
    except Exception as e:
        return f"Erro durante o processamento: {e}"

def iniciar_interface():
    def selecionar_excel():
        caminho = filedialog.askopenfilename(filetypes=[("Arquivos Excel", "*.xlsx")])
        entry_excel.delete(0, tk.END)
        entry_excel.insert(0, caminho)

    def selecionar_word_base():
        caminho = filedialog.askopenfilename(filetypes=[("Documentos Word", "*.docx")])
        entry_word_base.delete(0, tk.END)
        entry_word_base.insert(0, caminho)

    def selecionar_saida():
        caminho = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos Word", "*.docx")])
        entry_saida.delete(0, tk.END)
        entry_saida.insert(0, caminho)

    def atualizar_progresso(valor):
        barra_progresso["value"] = valor
        root.update_idletasks()

    def executar():
        def tarefa():
            barra_progresso["value"] = 0
            excel = entry_excel.get()
            word_entrada = entry_word_base.get()
            word_saida = entry_saida.get()

            if not (os.path.isfile(excel) and os.path.isfile(word_entrada)):
                messagebox.showerror("Erro", "Verifique se os arquivos de entrada existem.")
                return
            
            resultado = processar(excel, word_entrada, word_saida, progresso_callback=atualizar_progresso)
            barra_progresso["value"] = 100
            messagebox.showinfo("Resultado", resultado)

        threading.Thread(target=tarefa).start()

    root = tk.Tk()
    root.title("Atualizador de Documento Word - Consulta Pública 7217")

    tk.Label(root, text="Arquivo Excel:").grid(row=0, column=0, sticky="e")
    entry_excel = tk.Entry(root, width=60)
    entry_excel.grid(row=0, column=1)
    tk.Button(root, text="Selecionar", command=selecionar_excel).grid(row=0, column=2)

    tk.Label(root, text="Word Base (.docx):").grid(row=1, column=0, sticky="e")
    entry_word_base = tk.Entry(root, width=60)
    entry_word_base.grid(row=1, column=1)
    tk.Button(root, text="Selecionar", command=selecionar_word_base).grid(row=1, column=2)

    tk.Label(root, text="Salvar Word Final como:").grid(row=2, column=0, sticky="e")
    entry_saida = tk.Entry(root, width=60)
    entry_saida.grid(row=2, column=1)
    tk.Button(root, text="Selecionar", command=selecionar_saida).grid(row=2, column=2)

    tk.Button(root, text="Executar", command=executar, bg="green", fg="white").grid(row=3, column=1, pady=10)

    # Barra de progresso
    barra_progresso = ttk.Progressbar(root, orient="horizontal", length=400, mode="determinate")
    barra_progresso.grid(row=4, column=0, columnspan=3, pady=10)

    root.mainloop()

if __name__ == "__main__":
    iniciar_interface()
