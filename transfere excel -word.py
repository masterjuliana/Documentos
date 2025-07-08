import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

# Definir caminhos dos arquivos
arquivo_excel = r"W:\MINISTÉRIO DAS CIDADES\Consulta 7217\python\Consideracoes-sobre-a-Consulta_Publica_Decreto_7217.2010.xlsx"
arquivo_word_saida = r"W:\MINISTÉRIO DAS CIDADES\Consulta 7217\python\saida.docx"

try:
    df = pd.read_excel(arquivo_excel, engine="openpyxl")
except FileNotFoundError:
    raise FileNotFoundError(f"O arquivo '{arquivo_excel}' não foi encontrado.")
except Exception as e:
    raise RuntimeError(f"Erro ao ler o arquivo Excel: {e}")

# Remover espaços extras dos nomes das colunas
df.columns = df.columns.str.strip()

# Colunas desejadas
colunas_desejadas = ['Item CP alterado', 'Numero', 'Titulo da Contribuição ', 'Texto', 'Justificativa', 'Nome']

# Verificar se todas as colunas existem
colunas_ausentes = [col for col in colunas_desejadas if col not in df.columns]
if colunas_ausentes:
    raise ValueError(f"Colunas ausentes no Excel: {', '.join(colunas_ausentes)}")

# Selecionar as colunas desejadas
df = df[colunas_desejadas]

if df.empty:
    raise ValueError("Nenhum dado encontrado após a seleção das colunas! Verifique o conteúdo do arquivo Excel.")

# Substituir _x000D_ por aspas "
df = df.replace(r'_x000D_', '"', regex=True)

# Converter "Item CP alterado" para números e filtrar a partir do item 100
df["Item CP alterado"] = pd.to_numeric(df["Item CP alterado"], errors="coerce")
df = df.dropna(subset=["Item CP alterado"])
df["Item CP alterado"] = df["Item CP alterado"].astype(int)
df = df[df["Item CP alterado"] >= 100]

# Agrupar contribuições corretamente
df_contribuicoes_agrupado = df.groupby("Item CP alterado").agg({
    "Numero": lambda x: ", ".join(x.astype(str)),
    "Texto": lambda x: "\n".join(x),
    "Justificativa": lambda x: "\n".join(x),
    "Nome": lambda x: ", ".join(x)
}).reset_index()

# Abrir o documento Word
doc = Document(arquivo_word_saida)

# Iterar pelas tabelas do documento para encontrar o local correto
for table in doc.tables:
    for row in table.rows:
        item_cp_cell = row.cells[0].text.strip()  # Coluna "Item" no Word

        # Verificar se "Item" do Word corresponde ao "Item CP alterado"
        if item_cp_cell.isdigit():
            item_cp = int(item_cp_cell)

            # Encontrar as contribuições associadas
            matching_row = df_contribuicoes_agrupado[df_contribuicoes_agrupado["Item CP alterado"] == item_cp]

            if not matching_row.empty:
                # Adicionar "Item CP alterado"
                row.cells[0].text = str(matching_row["Item CP alterado"].values[0])
                p_item = row.cells[0].paragraphs[0].runs[0]
                p_item.bold = True
                p_item.font.name = "Calibri"
                p_item.font.size = Pt(10)

                # Adicionar "Numero" com negrito e tamanho 8
                row.cells[1].text = str(matching_row["Numero"].values[0])
                p_numero = row.cells[1].paragraphs[0].runs[0]
                p_numero.bold = True
                p_numero.font.name = "Calibri"
                p_numero.font.size = Pt(8)

                # Adicionar "Visões das contribuições" com Texto, Justificativa e Nome corretamente formatados
                texto_completo = f"{matching_row['Texto'].values[0]}\n{matching_row['Justificativa'].values[0]}\n{matching_row['Nome'].values[0]}"
                row.cells[2].text = texto_completo

                # Formatar fonte da coluna "Visões das contribuições" corretamente
                p_visoes = row.cells[2].paragraphs[0].runs[0]
                p_visoes.font.name = "Calibri"
                p_visoes.font.size = Pt(8)

# Salvar o arquivo ajustado
arquivo_word_ajustado = r"W:\MINISTÉRIO DAS CIDADES\Consulta 7217\python\saida_ajustada.docx"
doc.save(arquivo_word_ajustado)

print(f"Arquivo Word ajustado criado: {arquivo_word_ajustado}")
