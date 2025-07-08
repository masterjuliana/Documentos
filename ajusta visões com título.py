import pandas as pd
import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from collections import defaultdict
import re

# ==============================================================================
# 1. CONFIGURAÇÃO INICIAL
# ==============================================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("processamento.log"),
        logging.StreamHandler()
    ]
)

BASE_DIR = Path(__file__).parent
ARQUIVO_EXCEL = BASE_DIR / "Consideracoes-sobre-a-Consulta_Publica_Decreto_7217.2010.xlsx"
ARQUIVO_WORD_ENTRADA = BASE_DIR / "saida.docx"
ARQUIVO_WORD_SAIDA = BASE_DIR / "saida_título.docx"

# ==============================================================================
# 2. FUNÇÕES AUXILIARES
# ==============================================================================

def limpar_celula(celula) -> None:
    """Remove todo o conteúdo de uma célula preservando formatação básica"""
    for paragraph in celula.paragraphs:
        for run in paragraph.runs:
            run.text = ""
    while len(celula.paragraphs) >= 52:
        celula._element.remove(celula.paragraphs[0]._element)

def adicionar_linha_separadora(paragraph):
    """Adiciona uma borda inferior ao parágrafo para criar uma linha separadora"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    p_pr.append(p_bdr)
    
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), 'auto')
    
    p_bdr.append(bottom_border)

def formatar_celula_com_contribuicoes(celula, contribuicoes: list, fonte: str = 'Calibri', tamanho: int = 8):
    """Preenche a célula com as contribuições usando linhas gráficas como separadores"""
    limpar_celula(celula)
    
    for idx, contrib in enumerate(contribuicoes):
        # Adicionar número da contribuição
        p_num = celula.add_paragraph()
        run_num = p_num.add_run(contrib['numero'])
        run_num.font.name = fonte
        run_num.font.size = Pt(tamanho)
        run_num.bold = True
        
        # Adicionar Titulo da Contribuição 
        p_num = celula.add_paragraph()
        run_num = p_num.add_run(contrib['Titulo da Contribuição'])
        run_num.font.name = fonte
        run_num.font.size = Pt(tamanho)
        
        # Adicionar texto da contribuição
        p_texto = celula.add_paragraph(contrib['texto'])
        for run in p_texto.runs:
            run.font.name = fonte
            run.font.size = Pt(tamanho)
        
        # Adicionar nome do autor
        p_autor = celula.add_paragraph()
        run_autor = p_autor.add_run(f"({contrib['nome']})")
        run_autor.font.name = fonte
        run_autor.font.size = Pt(tamanho)
        run_autor.italic = True
        
        # Adicionar linha separadora (exceto após a última contribuição)
        if idx < len(contribuicoes) - 1:
            p_separador = celula.add_paragraph()
            adicionar_linha_separadora(p_separador)

def validar_item(item: str) -> int:
    """Valida e converte valores de item com tratamento robusto"""
    try:
        item_limpo = re.sub(r'\D', '', str(item))
        return int(item_limpo) if item_limpo else None
    except (ValueError, TypeError):
        return None

# ==============================================================================
# 3. FUNÇÃO PRINCIPAL DE EXECUÇÃO
# ==============================================================================

def main() -> None:
    # Leitura e preparação dos dados
    logging.info(f"Iniciando leitura do Excel: {ARQUIVO_EXCEL}")
    
    try:
        df = pd.read_excel(
            ARQUIVO_EXCEL,
            engine="openpyxl",
            dtype=str,
            usecols=['Item CP alterado', 'Numero', 'Titulo da Contribuição ','Texto', 'Nome']
            )

        logging.info(f"Colunas lidas do Excel: {df.columns.tolist()}") # Adicione esta linha temporariamente
    except Exception as e:
        logging.exception(f"Falha na leitura do Excel: {e}")
        return

    # Pré-processamento de dados
    for col in df.columns:
        if df[col].dtype == 'object': # Verifica se a coluna contém strings
            df[col] = df[col].astype(str).str.replace('_x000D_', '\n', regex=False)
    df.columns = df.columns.str.strip()
    
    # Filtragem e validação de itens
    df['Item_validado'] = df['Item CP alterado'].apply(validar_item)
    df = df.dropna(subset=['Item_validado'])
    df['Item_validado'] = df['Item_validado'].astype(int)
    df = df[df['Item_validado'] == 0]
    
    if df.empty:
        logging.warning("Nenhum dado válido encontrado após filtragem!")
        return

    # Agrupamento de contribuições em formato estruturado
    contrib_por_item = defaultdict(list)
    df = df.sort_values(by=["Item_validado", "Numero"])
    
    for _, row in df.iterrows():
        item = row['Item_validado']
        contrib = {
            'numero': f"{row['Numero']}" if pd.notna(row['Numero']) else "[sem número]",
            'Titulo da Contribuição': row['Titulo da Contribuição'] if pd.notna(row['Titulo da Contribuição']) else "[sem Titulo da Contribuição]", 
            'texto': row['Texto'] if pd.notna(row['Texto']) else "[sem texto]",
            'nome': row['Nome'] if pd.notna(row['Nome']) else "[autor desconhecido]"
        }
        contrib_por_item[item].append(contrib)
    
    logging.info(f"Contribuições agrupadas para {len(contrib_por_item)} itens")

    # Processamento do documento Word
    try:
        doc = Document(ARQUIVO_WORD_ENTRADA)
    except Exception as e:
        logging.exception(f"Falha ao abrir documento Word: {e}")
        return

    itens_atualizados = 0
    
    for tabela in doc.tables:
        for linha in tabela.rows:
            if len(linha.cells) < 3:
                continue
                
            celula_item = linha.cells[0]
            item_cp = validar_item(celula_item.text)
            
            if item_cp and item_cp in contrib_por_item:
                formatar_celula_com_contribuicoes(
                    linha.cells[2], 
                    contrib_por_item[item_cp],
                    tamanho=8
                )
                itens_atualizados += 1

    # Salvamento do resultado
    try:
        doc.save(ARQUIVO_WORD_SAIDA)
        logging.info(f"Documento salvo com sucesso: {ARQUIVO_WORD_SAIDA}")
        logging.info(f"Total de itens atualizados: {itens_atualizados}")
    except Exception as e:
        logging.exception(f"Erro ao salvar documento: {e}")

# ==============================================================================
# 4. EXECUÇÃO
# ==============================================================================

if __name__ == "__main__":
    main()