import sys
import logging
import pandas as pd
import unicodedata
from docx import Document
from docx.shared import Pt
from PyQt5.QtWidgets import (
    QApplication, QWidget, QFileDialog, QPushButton, QLabel,
    QVBoxLayout, QHBoxLayout, QLineEdit, QCheckBox, QMessageBox
)

# ==============================================================================
# FUNÇÕES AUXILIARES
# ==============================================================================

def limpar_texto(texto):
    if pd.isna(texto):
        return ""
    return unicodedata.normalize("NFKC", str(texto).strip())


def formatar_celula(celula, texto, nome_fonte='Calibri', tamanho_fonte=10, negrito=False):
    celula.text = texto
    paragraph = celula.paragraphs[0]
    run = paragraph.add_run() if not paragraph.runs else paragraph.runs[0]
    run.font.name = nome_fonte
    run.font.size = Pt(tamanho_fonte)
    run.bold = negrito


def processar_contribuicoes(arquivo_excel, word_entrada, word_saida, debug=False):
    if debug:
        logging.basicConfig(level=logging.DEBUG)
    else:
        logging.basicConfig(level=logging.INFO)

    logging.info(f"Lendo Excel: {arquivo_excel}")
    df = pd.read_excel(arquivo_excel, engine="openpyxl")
    df.columns = df.columns.str.strip()

    colunas_desejadas = ['Item CP alterado', 'Numero', 'Texto', 'Justificativa', 'Nome']
    df = df[colunas_desejadas].dropna()
    df["Item CP alterado"] = pd.to_numeric(df["Item CP alterado"], errors="coerce").astype('Int64')
    df = df.dropna(subset=["Item CP alterado"])
    df = df[df["Item CP alterado"] >= 100]

    df_agrupado = df.groupby("Item CP alterado").agg(lambda x: list(x)).reset_index()

    logging.info(f"Abrindo documento Word: {word_entrada}")
    doc = Document(word_entrada)

    itens_atualizados = 0
    for table in doc.tables:
        for row in table.rows:
            item_cp_cell_text = row.cells[0].text.strip()
            if item_cp_cell_text.isdigit():
                item_cp_word = int(item_cp_cell_text)
                matching_row = df_agrupado[df_agrupado["Item CP alterado"] == item_cp_word]

                if not matching_row.empty:
                    contrib_list = []
                    for num, texto, justificativa, nome in zip(
                        matching_row["Numero"].values[0],
                        matching_row["Texto"].values[0],
                        matching_row["Justificativa"].values[0],
                        matching_row["Nome"].values[0]
                    ):
                        if pd.notna(texto) and pd.notna(justificativa):
                            contrib_list.append(
                                f"CP-{item_cp_word}: {limpar_texto(num)}\r\n"
                                f"{limpar_texto(texto)}\r\n"
                                f"{limpar_texto(justificativa)}\r\n"
                                f"({limpar_texto(nome)})"
                            )

                    if contrib_list:
                        texto_completo = "\r\n\r\n".join(contrib_list)
                        logging.debug(f"Item {item_cp_word}, Visões:\n{texto_completo}")
                        formatar_celula(row.cells[2], texto=texto_completo, tamanho_fonte=8)

                    formatar_celula(row.cells[0], texto=str(item_cp_word), tamanho_fonte=10, negrito=True)
                    formatar_celula(row.cells[1], texto="")
                    itens_atualizados += 1

    doc.save(word_saida)
    logging.info(f"{itens_atualizados} itens atualizados. Word salvo em {word_saida}")


# ==============================================================================
# INTERFACE PyQt5
# ==============================================================================

class App(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Atualizador de Word - Consulta Pública")
        self.setGeometry(100, 100, 600, 300)

        layout = QVBoxLayout()

        self.excel_input = self.criar_linha_input("Arquivo Excel:")
        self.word_input = self.criar_linha_input("Word de entrada:")
        self.word_output = self.criar_linha_input("Word de saída:")

        self.debug_checkbox = QCheckBox("Ativar modo debug")
        layout.addWidget(self.debug_checkbox)

        self.btn_executar = QPushButton("Executar")
        self.btn_executar.clicked.connect(self.executar)
        layout.addWidget(self.btn_executar)

        self.status = QLabel("")
        layout.addWidget(self.status)

        self.setLayout(layout)

    def criar_linha_input(self, label_text):
        layout = QHBoxLayout()
        label = QLabel(label_text)
        entrada = QLineEdit()
        botao = QPushButton("Selecionar")
        botao.clicked.connect(lambda: self.abrir_arquivo(entrada))
        layout.addWidget(label)
        layout.addWidget(entrada)
        layout.addWidget(botao)
        self.layout().addLayout(layout) if self.layout() else None
        return entrada

    def abrir_arquivo(self, line_edit):
        file_name, _ = QFileDialog.getOpenFileName(self, "Selecionar arquivo")
        if file_name:
            line_edit.setText(file_name)

    def executar(self):
        excel_path = self.excel_input.text()
        word_entrada = self.word_input.text()
        word_saida = self.word_output.text()
        debug = self.debug_checkbox.isChecked()

        if not all([excel_path, word_entrada, word_saida]):
            QMessageBox.warning(self, "Erro", "Por favor, selecione todos os arquivos.")
            return

        try:
            processar_contribuicoes(excel_path, word_entrada, word_saida, debug)
            self.status.setText("✅ Processamento concluído com sucesso!")
        except Exception as e:
            self.status.setText("❌ Erro no processamento.")
            QMessageBox.critical(self, "Erro", str(e))


# ==============================================================================
# MAIN
# ==============================================================================

if __name__ == "__main__":
    app = QApplication(sys.argv)
    janela = App()
    janela.show()
    sys.exit(app.exec_())
